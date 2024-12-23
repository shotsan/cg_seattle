import streamlit as st
import PyPDF2
import docx
from docx.shared import Pt
import tempfile
import os
from pdf2image import convert_from_path
import io

st.set_page_config(layout="wide")


def year_to_words(year):
    ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
    tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
    teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
    
    def convert_two_digits(n):
        if n < 10:
            return ones[n]
        elif n < 20:
            return teens[n-10]
        else:
            return tens[n//10] + (' ' + ones[n%10] if n%10 != 0 else '')
    
    if not isinstance(year, int):
        year = int(year)
    
    first_two = year // 100
    last_two = year % 100
    
    result = convert_two_digits(first_two)
    if last_two:
        result += ' ' + convert_two_digits(last_two)
        
    return result

def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_fields_from_pdf(text):
    fields = {
        'title': '',
        'given_name': '',
        'surname': '',
        'passport_no': '',
        'date_of_birth': '',
        'place_of_birth': '',
        'state_of_birth': '',
        'date_of_issue': '',
        'place_of_issue': '',
        'father_name': '',
        'mother_name': '',
        'relation': '',
    }
    
    lines = text.split('\n')
    for line in lines:
        if 'TITLE :' in line:
            fields['title'] = line.split(':')[1].strip()
        elif 'GIVEN NAME (AS PER PASSPORT) :' in line:
            fields['given_name'] = line.split(':')[1].strip()
        elif 'SURNAME (AS PER PASSPORT) :' in line:
            fields['surname'] = line.split(':')[1].strip()
        elif 'PASSPORT NO :' in line:
            fields['passport_no'] = line.split(':')[1].strip()
        elif 'DATE OF BIRTH :' in line:
            fields['date_of_birth'] = line.split(':')[1].strip()
        elif 'PLACE OF BIRTH :' in line:
            fields['place_of_birth'] = line.split(':')[1].strip()
        elif 'STATE OF BIRTH :' in line:
            fields['state_of_birth'] = line.split(':')[1].strip()
        elif 'MOTHER NAME :' in line:
            fields['mother_name'] = line.split(':')[1].strip()
        elif 'FATHER NAME :' in line:
            fields['father_name'] = line.split(':')[1].strip()
        elif 'DATE OF ISSUE :' in line:
            fields['date_of_issue'] = line.split(':')[1].strip()
        elif 'PLACE OF ISSUE :' in line:
            fields['place_of_issue'] = line.split(':')[1].strip()
        if fields['title'] =='MR.':
            fields['relation'] = 'S/o'
        elif fields['title'] =='MRS.':
            fields['relation'] = 'D/o'
    
    return fields

def create_word_doc(fields, template_path):
    doc = docx.Document(template_path)
    
    template_vars = {
        '${TITLE}': fields['title'],
        '${GIVEN_NAME}': fields['given_name'],
        '${SURNAME}': fields['surname'],
        '${PASSPORT_NO}': fields['passport_no'],
        '${DATE_OF_BIRTH}': fields['date_of_birth'],
        '${PLACE_OF_BIRTH}': fields['place_of_birth'],
        '${STATE_OF_BIRTH}': fields['state_of_birth'],
        '${MOTHER_NAME}': fields['mother_name'],
        '${FATHER_NAME}': fields['father_name'],
        '${DATE_OF_ISSUE}': fields['date_of_issue'],
        '${PLACE_OF_ISSUE}': fields['place_of_issue'],
        '${RELATION}': fields['relation'],
        '${BIRTH_YEAR_IN_WORDS}': year_to_words(fields['date_of_birth'].split(' ')[-1]),
        '${TITLE1}': fields['title'],
    }
    
    def replace_text_preserve_formatting(paragraph):
        # Get the full paragraph text
        paragraph_text = paragraph.text
        has_replacement = False
        #st.write("Processing paragraph:", paragraph.text)

        # Check if any replacements are needed
        for key in template_vars:
            if key in paragraph_text:
                #st.write(f"Found key {key} in paragraph")

                has_replacement = True
                break
        
        if has_replacement:
            # Store runs and their properties
            runs_with_formatting = []
            for run in paragraph.runs:
                current_text = run.text
                for key, value in template_vars.items():
                    if key in current_text:
                        current_text = current_text.replace(key, value)
                # Preserve bold and other formatting
                runs_with_formatting.append((current_text, run.bold, run.italic, run.font.name))
            
            # Clear paragraph and rebuild with formatting
            paragraph.clear()
            for text, bold, italic, font in runs_with_formatting:
                new_run = paragraph.add_run(text)
                new_run.bold = bold
                new_run.italic = italic
                if font:
                    new_run.font.name = font

    # Process main document paragraphs
    for paragraph in doc.paragraphs:
        replace_text_preserve_formatting(paragraph)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserve_formatting(paragraph)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer
def preview_generated_doc(doc_buffer):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_file.write(doc_buffer.getvalue())
        doc_path = tmp_file.name
    
    preview_doc = docx.Document(doc_path)
    preview_text = []
    
    for paragraph in preview_doc.paragraphs:
        preview_text.append(paragraph.text)
    
    for table in preview_doc.tables:
        for row in table.rows:
            row_text = [cell.text.ljust(20) for cell in row.cells]
            preview_text.append(' | '.join(row_text))
    
    os.unlink(doc_path)
    

    return '\n'.join(preview_text)

def main():
    with st.container():
        st.markdown("""
            <style>
                .title {
                    text-align: center !important;
                    font-size: 80px !important;
                    font-weight: 900 !important;
                    padding: 25px;
                    background: linear-gradient(45deg, #FF9933, #FFFFFF, #138808);
                    -webkit-background-clip: text;
                    -webkit-text-fill-color: transparent;
                    font-family: 'Arial', sans-serif;
                    text-transform: uppercase;
                    margin-bottom: 30px;
                    text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.1);
                    line-height: 1.2;
                    display: block;
                }
                /* Override any potential Streamlit default styles */
                .stMarkdown {
                    max-width: 100% !important;
                }
            </style>
        """, unsafe_allow_html=True)

        st.markdown('<h1 class="title">Automated Application Processing Platform</h1>', unsafe_allow_html=True)





        
        template_paths = {
            "Birth Template": "templates/birth_template.docx",
        }
        
        col1, col2 = st.columns([4, 6])
        
        with col1:
            st.header("Upload Application")
            uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
            
            template_option = st.selectbox(
                "Choose Template",
                list(template_paths.keys())
            )
            
            if uploaded_file:
                text = extract_text_from_pdf(uploaded_file)
                fields = extract_fields_from_pdf(text)
                
                st.subheader("Applicant Information")
                st.json(fields)
            
                st.markdown("""
                    <style>
                        .stButton > button {
                            background-color: #000080;
                            color: white;
                            font-weight: 900;
                            font-size: 24px;
                            padding: 15px 30px;
                            border-radius: 5px;
                            border: none;
                            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                            transition: all 0.3s ease;
                            text-transform: uppercase;
                            letter-spacing: 1px;
                        }
                        .stButton > button:hover {
                            background-color: #000066;
                            transform: translateY(-2px);
                            box-shadow: 0 6px 8px rgba(0, 0, 0, 0.2);
                        }
                    </style>
                """, unsafe_allow_html=True)











            if st.button("Generate Certificate"):
                doc_buffer = create_word_doc(fields, template_paths[template_option])
    
    # Show preview
                st.subheader("Certificate Preview")
                preview_text = preview_generated_doc(doc_buffer)
                st.text_area("Preview", preview_text, height=300)
    
    # Download button
                st.download_button(
                 label="Download Certificate",
                data=doc_buffer,
                file_name="filled_template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                 )
        
        with col2:
            st.header("Application Preview")
            if uploaded_file:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    pdf_path = tmp_file.name
                
                images = convert_from_path(pdf_path)
                for image in images:
                    st.image(image, use_container_width=True)
                
                os.unlink(pdf_path)

if __name__ == "__main__":
    main()
