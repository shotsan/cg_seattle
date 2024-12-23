from docx import Document
from docx.shared import Pt, Inches

def create_basic_template():
    doc = Document()
    doc.add_heading('Document Title', 0)
    doc.add_paragraph('${content}')
    doc.save('templates/basic.docx')

def create_business_template():
    doc = Document()
    doc.add_heading('Business Letter', 0)
    doc.add_paragraph('Date: [Current Date]')
    doc.add_paragraph('Dear Sir/Madam,')
    doc.add_paragraph('${content}')
    doc.add_paragraph('Best Regards,')
    doc.add_paragraph('[Your Name]')
    doc.save('templates/business.docx')

def create_report_template():
    doc = Document()
    doc.add_heading('Report', 0)
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph('${content}')
    doc.add_heading('Findings', 1)
    doc.add_paragraph('[Findings will be inserted here]')
    doc.save('templates/report.docx')

if __name__ == "__main__":
    os.makedirs('templates', exist_ok=True)
    create_basic_template()
    create_business_template()
    create_report_template()

